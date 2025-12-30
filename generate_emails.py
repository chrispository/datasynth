import uuid
import random
import datetime
import os
import email
import re
from email.message import EmailMessage
from email.utils import formatdate
import mimetypes
from faker import Faker
from fpdf import FPDF
from docx import Document
from roster import RosterGenerator
from llm import GeminiGenerator
from pdf_utils import init_pdf

fake = Faker()


class FileGenerator:
    def __init__(self, output_dir="output", llm=None, topic=None):
        self.output_dir = output_dir
        self.llm = llm
        self.topic = topic
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)


    def create_pdf(self, filename, content_text):
        pdf = init_pdf()
        pdf.add_page()
        pdf.set_font("DejaVu", size=12)
        pdf.multi_cell(0, 10, txt=content_text)

        filepath = os.path.join(self.output_dir, filename)
        pdf.output(filepath)
        return filepath

    def create_docx(self, filename, content_text):
        doc = Document()
        # Add a title at the top of the Word document
        title_text = filename.replace(".docx", "").replace("_", " ")
        # Clean up common prefixes if any
        title_text = re.sub(r"^\d{4}\s+", "", title_text)
        doc.add_heading(title_text, 0)

        # Split by double newlines to create actual paragraphs
        paragraphs = content_text.split("\n\n")
        for p in paragraphs:
            if p.strip():
                doc.add_paragraph(p.strip())

        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def _generate_content(self, doc_type, context=None):
        """Generate document content using LLM or fallback templates."""
        if self.llm:
            prompt = f"Generate a realistic {doc_type} document"
            if self.topic:
                prompt += f" related to {self.topic}"
            if context:
                prompt += f". Context from related email thread: {context}"
            prompt += """. 

IMPORTANT: This is a standalone business document (Word/PDF attachment), NOT an email. 
Do not include any email language like greetings, signatures, "Dear", "Best regards", "From:", "To:", etc.
Make sure to use appropriate headings and sections where needed.
Keep it under 750 words. Write only the document content."""

            content = self.llm.generate_email_content(prompt)
            if content:
                return content

        # Fallback to topic-based template
        if self.topic:
            templates = [
                f"DOCUMENT: {self.topic}\n\n" + fake.paragraph(nb_sentences=8),
                f"REPORT: Analysis of {self.topic}\n\nExecutive Summary:\n"
                + fake.paragraph(nb_sentences=5)
                + "\n\nDetails:\n"
                + fake.paragraph(nb_sentences=8),
                f"NOTES: {self.topic} Discussion\n\n" + fake.paragraph(nb_sentences=10),
                f"PROPOSAL: {self.topic}\n\nBackground:\n"
                + fake.paragraph(nb_sentences=4)
                + "\n\nRecommendations:\n"
                + fake.paragraph(nb_sentences=6),
            ]
            return random.choice(templates)

        return fake.text(max_nb_chars=1000)

    def generate_random_file(self, base_name, doc_type="document", context=None):
        ext = random.choice(["pdf", "docx"])
        filename = f"{base_name}.{ext}"
        content = self._generate_content(doc_type, context)

        if ext == "pdf":
            return self.create_pdf(filename, content)
        else:
            return self.create_docx(filename, content)


class Attachment:
    def __init__(self, filename, filepath, content_type):
        self.id = str(uuid.uuid4())
        self.filename = filename
        self.filepath = filepath
        self.content_type = content_type

    def __repr__(self):
        return f"<Attachment {self.filename}>"


class Email:
    def __init__(
        self,
        sender,
        recipients,
        subject,
        body,
        date,
        message_id=None,
        parent_id=None,
        thread_id=None,
        msg_type="new",
    ):
        self.id = str(uuid.uuid4())
        self.message_id = (
            message_id if message_id else f"<{self.id}@{fake.free_email_domain()}>"
        )
        self.thread_id = thread_id if thread_id else str(uuid.uuid4())
        self.parent_id = (
            parent_id  # Points to the Email object or ID this is a reply/forward to
        )

        self.sender = sender
        self.recipients = recipients  # List of strings
        self.cc = []
        self.bcc = []

        self.subject = subject
        self.body = body
        self.date = date  # datetime object

        self.type = msg_type  # new, reply, forward

        self.in_reply_to = parent_id if parent_id else None
        self.references = []  # List of message_ids

        self.attachments = []

    def add_attachment(self, attachment):
        self.attachments.append(attachment)

    def __repr__(self):
        return f"[{self.date.strftime('%Y-%m-%d %H:%M')}] {self.sender} -> {', '.join(self.recipients)} | {self.subject} ({self.type})"


class ThreadGenerator:
    def __init__(
        self,
        roster=None,
        llm=None,
        start_date=None,
        output_dir="output",
        topic=None,
        attachment_percent=30,
    ):
        self.emails = []  # Flat list of all emails generated
        self.threads = {}  # Map thread_id -> list of Email objects
        self.current_date = (
            start_date
            if start_date
            else datetime.datetime.now() - datetime.timedelta(days=30)
        )
        self.roster = (
            roster
            if roster
            else [
                {
                    "name": fake.name(),
                    "email": fake.email(),
                    "title": "Employee",
                    "department": "General",
                }
                for _ in range(10)
            ]
        )
        self.llm = llm
        self.file_gen = FileGenerator(output_dir, llm=llm, topic=topic)
        self.topic = topic
        self.attachment_percent = attachment_percent / 100.0  # Convert to 0-1 range

    def _tick_time(self):
        # Advance time by random minutes/hours
        increment = datetime.timedelta(minutes=random.randint(1, 120))
        self.current_date += increment
        return self.current_date

    def get_person_display(self, person):
        return f"{person['name']} <{person['email']}>"

    def _get_thread_participants(self, thread_id):
        """Get all unique participants (senders and recipients) from a thread."""
        participants = set()
        if thread_id in self.threads:
            for email in self.threads[thread_id]:
                participants.add(email.sender)
                participants.update(email.recipients)
        return list(participants)

    def _can_forward_to_new_recipients(self, thread_id):
        """Check if there are roster members not in the current thread."""
        thread_participants = self._get_thread_participants(thread_id)
        participant_emails = set()
        for p in thread_participants:
            if " <" in p:
                participant_emails.add(p.split(" <")[1].rstrip(">"))
            else:
                participant_emails.add(p)

        for person in self.roster:
            if person["email"] not in participant_emails:
                return True
        return False

    def create_root_email(self):
        sender = random.choice(self.roster)
        recipients = random.sample(
            [p for p in self.roster if p != sender], k=random.randint(1, 3)
        )

        subject = None
        body = None

        if self.llm:
            subject, body = self.llm.generate_email(
                sender, recipients, self.topic if self.topic else "General check-in"
            )

        if not body:
            if self.topic:
                subject_start = random.choice(
                    [
                        "Regarding",
                        "Update on",
                        "Question about",
                        "Notes for",
                        "Discussion:",
                    ]
                )
                subject = f"{subject_start} {self.topic}"
                body_start = f"Hi all,\n\nI wanted to discuss {self.topic}."
                body = f"{body_start}\n\n" + fake.paragraph(nb_sentences=5)
            else:
                subject = fake.sentence(nb_words=4).rstrip(".")
                body = fake.paragraph(nb_sentences=5)

        email = Email(
            sender=self.get_person_display(sender),
            recipients=[self.get_person_display(r) for r in recipients],
            subject=subject,
            body=body,
            date=self._tick_time(),
            msg_type="new",
        )

        # Note: Attachments are now generated at save time for inclusive emails only

        self._store_email(email)
        return email

    def reply_to(self, parent_email, reply_all=True):
        # Extract name from display string "Name <email>"
        def parse_display(display):
            if " <" in display:
                name = display.split(" <")[0]
                email = display.split(" <")[1].rstrip(">")
                return {"name": name, "email": email}
            return {"name": display, "email": display}

        parent_recipients = [parse_display(r) for r in parent_email.recipients]
        parent_sender = parse_display(parent_email.sender)

        # Find the person in the roster that matches a recipient
        sender_info = random.choice(parent_recipients)
        # Try to find their full info in roster
        roster_sender = next(
            (p for p in self.roster if p["email"] == sender_info["email"]), None
        )
        if not roster_sender:
            roster_sender = {
                "name": sender_info["name"],
                "email": sender_info["email"],
                "title": "Employee",
                "department": "General",
            }

        # Primary recipient is always the original sender
        recipients = [parent_sender]
        roster_recipients = [
            next((p for p in self.roster if p["email"] == r["email"]), r)
            for r in recipients
        ]

        # Reply-all: CC other original recipients (excluding the new sender)
        cc_recipients = []
        if reply_all and len(parent_recipients) > 1:
            for r in parent_recipients:
                if r["email"] != sender_info["email"]:
                    cc_recipients.append(r)

        subject = parent_email.subject
        if not subject.lower().startswith("re:"):
            subject = f"Re: {subject}"

        new_body = None
        if self.llm:
            _, new_body = self.llm.generate_email(
                roster_sender, roster_recipients, self.topic, context=parent_email.body
            )

        if not new_body:
            new_body = fake.paragraph(nb_sentences=3)
            if self.topic and random.random() < 0.2:
                new_body += f"\n\nRegarding the {self.topic} aspect, I agree."

        # Recursively quote the ENTIRE parent body, indented
        # Clean up existing newlines before quoting
        parent_body_lines = parent_email.body.split("\n")
        quoted_lines = [f"> {line}" for line in parent_body_lines]
        quoted_block = "\n".join(quoted_lines)

        full_body = f"{new_body}\n\nOn {parent_email.date.strftime('%Y-%m-%d %H:%M')}, {parent_email.sender} wrote:\n{quoted_block}"

        email = Email(
            sender=self.get_person_display(roster_sender),
            recipients=[
                self.get_person_display(r)
                if isinstance(r, dict) and "email" in r
                else r
                for r in recipients
            ],
            subject=subject,
            body=full_body,
            date=self._tick_time(),
            message_id=None,
            parent_id=parent_email.message_id,
            thread_id=parent_email.thread_id,
            msg_type="reply",
        )

        # Add CC recipients for reply-all
        email.cc = [f"{r['name']} <{r['email']}>" for r in cc_recipients]

        # Handle Headers
        email.references = parent_email.references + [parent_email.message_id]

        self._store_email(email)
        return email

    def forward(self, parent_email):
        # Pick sender from thread participants for realism (someone who was in the conversation)
        thread_participants = self._get_thread_participants(parent_email.thread_id)

        # Find roster entries for participants
        def parse_display(display):
            if " <" in display:
                return display.split(" <")[1].rstrip(">")
            return display

        participant_emails = [parse_display(p) for p in thread_participants]
        roster_participants = [
            p for p in self.roster if p["email"] in participant_emails
        ]

        if roster_participants:
            sender = random.choice(roster_participants)
        else:
            # Fallback to random if no participants found in roster
            sender = random.choice(self.roster)

        # Forward to someone NOT in the thread
        potential_recipients = [
            p
            for p in self.roster
            if self.get_person_display(p) not in thread_participants
        ]
        if not potential_recipients:
            potential_recipients = [random.choice(self.roster)]
        recipients = [random.choice(potential_recipients)]

        subject = parent_email.subject
        if not subject.lower().startswith("fwd:"):
            subject = f"Fwd: {subject}"

        new_body = None
        if self.llm:
            _, new_body = self.llm.generate_email(
                sender,
                recipients,
                f"Forwarding: {self.topic if self.topic else parent_email.subject}",
                context=parent_email.body,
            )

        if not new_body:
            new_body = f"FYI."
            if self.topic and random.random() < 0.3:
                new_body = (
                    f"Thought you should see this regarding {self.topic}.\n\n"
                    + new_body
                )

        # Include the full parent body
        forward_block = f"---------- Forwarded message ----------\nFrom: {parent_email.sender}\nDate: {parent_email.date}\nSubject: {parent_email.subject}\nTo: {', '.join(parent_email.recipients)}\n\n{parent_email.body}"

        full_body = f"{new_body}\n\n{forward_block}"

        email = Email(
            sender=self.get_person_display(sender),
            recipients=[self.get_person_display(r) for r in recipients],
            subject=subject,
            body=full_body,
            date=self._tick_time(),
            message_id=None,
            parent_id=parent_email.message_id,
            thread_id=parent_email.thread_id,
            msg_type="forward",
        )

        email.references = parent_email.references + [parent_email.message_id]

        # Carry over attachments
        for att in parent_email.attachments:
            email.add_attachment(att)

        self._store_email(email)
        return email

    def _store_email(self, email):
        self.emails.append(email)
        if email.thread_id not in self.threads:
            self.threads[email.thread_id] = []
        self.threads[email.thread_id].append(email)

        inclusive_count = self._count_inclusive_emails()
        print(
            f"  [Progress] Total emails: {len(self.emails)} | Inclusive emails: {inclusive_count}",
            flush=True,
        )

    def _count_inclusive_emails(self):
        """Count emails that are not parents of other emails (leaf/inclusive emails)."""
        parent_message_ids = set()
        for email_obj in self.emails:
            if email_obj.parent_id:
                parent_message_ids.add(email_obj.parent_id)
        return sum(1 for e in self.emails if e.message_id not in parent_message_ids)

    def simulate(
        self, target_inclusive=5, max_emails_per_thread=9, early_end_chance=0.15
    ):
        """Simulate multiple email threads until we have target_inclusive leaf emails."""

        while self._count_inclusive_emails() < target_inclusive:
            # Create a new thread with a root email
            self.create_root_email()

            # Get the thread we just created
            tid = list(self.threads.keys())[-1]
            thread_msgs = self.threads[tid]

            # Determine this thread's target length (random between 2 and max_emails_per_thread)
            thread_target_length = random.randint(2, max_emails_per_thread)

            # Build this thread
            while (
                self._count_inclusive_emails() < target_inclusive
                and len(thread_msgs) < thread_target_length
            ):
                # Refresh thread_msgs reference (it may have been updated)
                thread_msgs = self.threads[tid]

                # Check if thread should end early (after at least 2 emails)
                if len(thread_msgs) >= 2 and random.random() < early_end_chance:
                    break

                # Pick a message to reply to - heavily bias towards most recent for realistic threading
                # 80% chance to reply to the most recent, 20% chance to reply to an older message (branching)
                parent = (
                    thread_msgs[-1]
                    if random.random() > 0.2
                    else random.choice(thread_msgs)
                )

                # Action weights: 80% reply, 10% forward, 10% skip (to add variation in timing)
                action = random.choices(
                    ["reply", "forward", "nothing"], weights=[0.8, 0.1, 0.1]
                )[0]

                # Convert forward to reply if no new recipients available (like syndata)
                if action == "forward" and not self._can_forward_to_new_recipients(tid):
                    action = "reply"

                if action == "reply":
                    self.reply_to(parent)
                elif action == "forward":
                    self.forward(parent)


def save_as_markdown(email_obj, output_dir="output", index=0):
    # Create safe subject for filename
    safe_subject = "".join([c if c.isalnum() else "_" for c in email_obj.subject])[:40]
    filename = f"{index:04d}_{safe_subject}.md"
    filepath = os.path.join(output_dir, filename)

    # Handle Attachments (move/rename them)
    att_list = []
    for att in email_obj.attachments:
        if os.path.exists(att.filepath):
            att_ext = os.path.splitext(att.filename)[1]
            att_base = os.path.splitext(att.filename)[0]
            new_att_name = f"{index:04d}_{safe_subject}_attachment_{att_base}{att_ext}"
            new_att_path = os.path.join(output_dir, new_att_name)

            if not os.path.exists(new_att_path) and os.path.exists(att.filepath):
                import shutil

                shutil.copy(att.filepath, new_att_path)

            att_list.append(new_att_name)
        else:
            att_list.append(att.filename)  # Just list name if file missing

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"**From:** {email_obj.sender}\n")
        f.write(f"**Date:** {email_obj.date.strftime('%A, %B %d, %Y %I:%M %p')}\n")
        f.write(f"**To:** {', '.join(email_obj.recipients)}\n")
        if email_obj.cc:
            f.write(f"**Cc:** {', '.join(email_obj.cc)}\n")
        f.write(f"**Subject:** {email_obj.subject}\n")

        if att_list:
            f.write(f"**Attachments:** {', '.join(att_list)}\n")

        f.write("\n---\n\n")
        f.write(email_obj.body)

    return filepath


if __name__ == "__main__":
    import argparse
    import sys
    import traceback
    from dotenv import load_dotenv

    load_dotenv()

    try:
        print("Starting generator...", flush=True)
        parser = argparse.ArgumentParser()
        parser.add_argument(
            "--files",
            type=int,
            default=5,
            help="Number of inclusive email threads to generate",
        )
        parser.add_argument(
            "--steps", type=int, default=None, help="Deprecated, use --files instead"
        )
        parser.add_argument(
            "--output", type=str, default="output", help="Output directory"
        )
        parser.add_argument("--topic", type=str, default=None, help="Topic to focus on")
        parser.add_argument(
            "--attachments",
            type=int,
            default=30,
            help="Percentage of emails with attachments (0-100)",
        )
        parser.add_argument(
            "--roster", type=str, default="roster.json", help="Path to roster file"
        )
        parser.add_argument(
            "--gemini", action="store_true", help="Use Gemini LLM for email generation"
        )
        parser.add_argument(
            "--model", type=str, default="gemini-1.5-flash", help="Gemini model to use"
        )
        # Kept for compatibility but ignored/removed
        parser.add_argument("--pdf", action="store_true", help="Ignored")
        args = parser.parse_args()

        # Handle Roster
        roster_gen = RosterGenerator()
        if os.path.exists(args.roster):
            print(f"Loading roster from {args.roster}...", flush=True)
            roster = roster_gen.load_roster(args.roster)
        else:
            print(f"Generating new roster...", flush=True)
            roster = roster_gen.generate_roster(25)
            roster_gen.save_roster(args.roster)
            print(f"Saved roster to {args.roster}", flush=True)

        # Handle LLM
        llm = None
        if args.gemini:
            if not os.getenv("GEMINI_API_KEY"):
                print("\nGemini API key not found.")
                key = input("Please paste your Gemini API key: ").strip()
                if key:
                    with open(".env", "a" if os.path.exists(".env") else "w") as f:
                        f.write(f"\nGEMINI_API_KEY={key}\n")
                    os.environ["GEMINI_API_KEY"] = key
                    print("API key saved to .env\n")
                else:
                    print(
                        "Error: Gemini API key is required for --gemini mode.",
                        file=sys.stderr,
                    )
                    sys.exit(1)

            print(f"Initializing Gemini LLM with model: {args.model}...", flush=True)
            llm = GeminiGenerator(model_name=args.model)

        # Create a topic-based subfolder for this run's output
        if args.topic:
            # Extract first two words from topic for folder name
            words = args.topic.split()[:2]
            folder_name = "_".join(w.lower() for w in words)
            # Clean up any non-alphanumeric chars
            folder_name = "".join(
                c if c.isalnum() or c == "_" else "_" for c in folder_name
            )
        else:
            # Fallback to timestamp if no topic
            folder_name = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        run_output_dir = os.path.join(args.output, folder_name)
        os.makedirs(run_output_dir, exist_ok=True)
        print(f"Output folder: {run_output_dir}", flush=True)

        gen = ThreadGenerator(
            roster=roster,
            llm=llm,
            output_dir=run_output_dir,
            topic=args.topic,
            attachment_percent=args.attachments,
        )

        # Handle backwards compatibility: --steps is deprecated in favor of --files
        target_files = args.files
        if args.steps is not None:
            print(f"Warning: --steps is deprecated, use --files instead", flush=True)
            target_files = args.steps

        print(f"Generating {target_files} inclusive email threads...", flush=True)
        print(f"Attachment rate: {args.attachments}%", flush=True)
        if args.topic:
            print(f"Topic: {args.topic}", flush=True)

        gen.simulate(target_inclusive=target_files)

        print(f"Generated {len(gen.emails)} emails.", flush=True)

        # Find all emails that are parents of other emails (non-inclusive)
        # Inclusive emails are "leaf" nodes - not referenced as parent_id by any other email
        parent_message_ids = set()
        for email_obj in gen.emails:
            if email_obj.parent_id:
                parent_message_ids.add(email_obj.parent_id)

        # Count inclusive emails
        inclusive_emails = [
            e for e in gen.emails if e.message_id not in parent_message_ids
        ]
        print(f"Inclusive (leaf) emails: {len(inclusive_emails)}", flush=True)

        # Sort inclusive emails by thread_id first, then by date
        # This ensures all emails from the same thread are grouped together
        inclusive_emails.sort(key=lambda e: (e.thread_id, e.date))

        # Only save inclusive emails (those not referenced as parents)
        # Attachments are generated at save time for inclusive emails only
        all_attachments = set()
        inclusive_idx = 0
        for email_obj in inclusive_emails:
                inclusive_idx += 1

                # Generate attachment for this inclusive email based on percentage
                if random.random() < args.attachments / 100.0:
                    safe_subject = "".join(
                        [c if c.isalnum() else "_" for c in email_obj.subject]
                    )[:40]
                    doc_types = ["report", "proposal", "notes", "analysis", "summary"]
                    doc_type = random.choice(doc_types)
                    filepath = gen.file_gen.generate_random_file(
                        safe_subject, doc_type=doc_type, context=email_obj.body[:200]
                    )
                    filename = os.path.basename(filepath)
                    ctype = (
                        "application/pdf"
                        if filename.endswith(".pdf")
                        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    email_obj.attachments = [Attachment(filename, filepath, ctype)]

                for att in email_obj.attachments:
                    all_attachments.add(att.filepath)

                md_path = save_as_markdown(
                    email_obj, gen.file_gen.output_dir, index=inclusive_idx
                )

                print(f"Saved: {md_path}", flush=True)

        # Cleanup original unnumbered attachment files
        for att_path in all_attachments:
            if os.path.exists(att_path):
                try:
                    os.remove(att_path)
                except Exception as e:
                    print(
                        f"Warning: Could not remove original attachment {att_path}: {e}"
                    )

    except Exception as e:
        print(f"\nCRITICAL ERROR: {e}", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)
