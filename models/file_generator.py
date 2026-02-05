"""
File generator for creating PDF and DOCX attachments.
"""

import os
import re
import random
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
from pdf_utils import init_pdf
from utils import sanitize_filename

from faker_instance import fake

STYLE_PROFILES = {
    "corporate_classic": {
        "body_font": "Calibri",
        "heading_font": "Cambria",
        "heading_color": RGBColor(0x1F, 0x3A, 0x5F),  # dark blue
        "body_size": Pt(11),
        "heading1_size": Pt(18),
        "heading2_size": Pt(14),
        "title_size": Pt(26),
        "line_spacing": 1.15,
        "space_after": Pt(8),
        "margins": Cm(2.54),  # 1.0"
        "table_style": "Medium Shading 1 Accent 1",
        "header_type": "company",
        "footer_type": "page_number",
    },
    "modern_minimal": {
        "body_font": "Arial",
        "heading_font": "Arial",
        "heading_color": RGBColor(0x22, 0x22, 0x22),  # near-black
        "body_size": Pt(10),
        "heading1_size": Pt(16),
        "heading2_size": Pt(13),
        "title_size": Pt(24),
        "line_spacing": 1.08,
        "space_after": Pt(6),
        "margins": Cm(1.91),  # 0.75"
        "table_style": "Light List Accent 1",
        "header_type": "none",
        "footer_type": "both",
    },
    "formal_report": {
        "body_font": "Times New Roman",
        "heading_font": "Georgia",
        "heading_color": RGBColor(0x8B, 0x00, 0x00),  # dark red
        "body_size": Pt(12),
        "heading1_size": Pt(18),
        "heading2_size": Pt(15),
        "title_size": Pt(28),
        "line_spacing": 1.5,
        "space_after": Pt(10),
        "margins": Cm(3.18),  # 1.25"
        "table_style": "Medium Grid 1 Accent 2",
        "header_type": "confidential",
        "footer_type": "date",
    },
    "tech_memo": {
        "body_font": "Consolas",
        "heading_font": "Calibri",
        "heading_color": RGBColor(0x00, 0x70, 0xC0),  # bright blue
        "body_size": Pt(10),
        "heading1_size": Pt(16),
        "heading2_size": Pt(13),
        "title_size": Pt(22),
        "line_spacing": 1.0,
        "space_after": Pt(4),
        "margins": Cm(2.03),  # 0.8"
        "table_style": "Light Grid Accent 5",
        "header_type": "draft",
        "footer_type": "page_number",
    },
    "executive_brief": {
        "body_font": "Garamond",
        "heading_font": "Garamond",
        "heading_color": RGBColor(0x1B, 0x5E, 0x20),  # dark green
        "body_size": Pt(12),
        "heading1_size": Pt(20),
        "heading2_size": Pt(16),
        "title_size": Pt(30),
        "line_spacing": 1.3,
        "space_after": Pt(12),
        "margins": Cm(3.81),  # 1.5"
        "table_style": "Medium Shading 2 Accent 3",
        "header_type": "company",
        "footer_type": "both",
    },
    "compact_dense": {
        "body_font": "Verdana",
        "heading_font": "Tahoma",
        "heading_color": RGBColor(0x60, 0x60, 0x60),  # medium gray
        "body_size": Pt(9),
        "heading1_size": Pt(14),
        "heading2_size": Pt(11),
        "title_size": Pt(20),
        "line_spacing": 1.0,
        "space_after": Pt(3),
        "margins": Cm(1.27),  # 0.5"
        "table_style": "Light List Accent 3",
        "header_type": "none",
        "footer_type": "none",
    },
}


class FileGenerator:
    """
    Generates PDF and DOCX file attachments.
    
    Uses LLM for content generation when available, with faker-based
    fallbacks for offline operation.
    """
    
    def __init__(self, output_dir: str = "output", llm: Optional[object] = None, topic: Optional[str] = None) -> None:
        self.output_dir = output_dir
        self.llm = llm
        self.topic = topic
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def create_pdf(self, filename: str, content_text: str) -> str:
        """Create a PDF file with the given content."""
        pdf = init_pdf()
        pdf.add_page()
        pdf.set_font("DejaVu", size=12)
        pdf.multi_cell(0, 10, txt=content_text)

        filepath = os.path.join(self.output_dir, filename)
        pdf.output(filepath)
        return filepath

    def create_docx(self, filename: str, content_text: str) -> str:
        """Create a DOCX file with rich formatting parsed from markdown."""
        doc = Document()

        # Pick a random style profile
        profile_name = random.choice(list(STYLE_PROFILES.keys()))
        profile = STYLE_PROFILES[profile_name]

        # Apply style profile (fonts, colors, margins, spacing)
        self._apply_style_profile(doc, profile)

        # 25% chance of landscape orientation
        if random.random() < 0.25:
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

        # Add header/footer per profile
        self._add_header_footer(doc, profile)

        # Add a title at the top of the Word document
        title_text = filename.replace(".docx", "").replace("_", " ")
        title_text = re.sub(r"^\d{4}\s+", "", title_text)
        doc.add_heading(title_text, 0)

        lines = content_text.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Headings: ## or ###
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip("# "), level=2)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip("# "), level=1)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:].strip("# "), level=1)
            # Bullet list: - item or * item
            elif re.match(r"^[-*]\s+", stripped):
                text = re.sub(r"^[-*]\s+", "", stripped)
                para = doc.add_paragraph(style="List Bullet")
                self._add_runs_with_bold(para, text)
            # Numbered list: 1. item
            elif re.match(r"^\d+\.\s+", stripped):
                text = re.sub(r"^\d+\.\s+", "", stripped)
                para = doc.add_paragraph(style="List Number")
                self._add_runs_with_bold(para, text)
            else:
                para = doc.add_paragraph()
                self._add_runs_with_bold(para, stripped)

        # Maybe add a data table
        self._maybe_add_table(doc, profile)

        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def _add_runs_with_bold(self, paragraph: object, text: str) -> None:
        """Parse **bold** markers in text and add runs to paragraph."""
        parts = re.split(r"(\*\*[^*]+\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def _apply_style_profile(self, doc: Document, profile: dict) -> None:
        """Apply a visual style profile to the document (fonts, colors, margins, spacing)."""
        style_normal = doc.styles["Normal"]
        style_normal.font.name = profile["body_font"]
        style_normal.font.size = profile["body_size"]
        pf = style_normal.paragraph_format
        pf.line_spacing = profile["line_spacing"]
        pf.space_after = profile["space_after"]

        for style_name, size_key in [("Title", "title_size"), ("Heading 1", "heading1_size"), ("Heading 2", "heading2_size")]:
            style = doc.styles[style_name]
            style.font.name = profile["heading_font"]
            style.font.size = profile[size_key]
            style.font.color.rgb = profile["heading_color"]

        section = doc.sections[0]
        m = profile["margins"]
        section.top_margin = m
        section.bottom_margin = m
        section.left_margin = m
        section.right_margin = m

    def _add_header_footer(self, doc: Document, profile: dict) -> None:
        """Add header and/or footer to the document based on the profile."""
        section = doc.sections[0]

        # Header
        header_type = profile["header_type"]
        if header_type != "none":
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0]
            if header_type == "company":
                p.text = fake.company()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            elif header_type == "confidential":
                p.text = "CONFIDENTIAL"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            elif header_type == "draft":
                p.text = "DRAFT"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Footer
        footer_type = profile["footer_type"]
        if footer_type != "none":
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0]
            if footer_type == "page_number":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_page_number_field(p, profile)
            elif footer_type == "date":
                p.text = fake.date_this_year().strftime("%B %d, %Y")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            elif footer_type == "both":
                p.text = fake.date_this_year().strftime("%B %d, %Y")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                # Add right-aligned tab stop with page number
                run_sep = p.add_run("\t\t")
                run_sep.font.size = Pt(8)
                self._add_page_number_field(p, profile)

    def _add_page_number_field(self, paragraph, profile: dict) -> None:
        """Insert a PAGE field code into the paragraph for automatic page numbering."""
        run = paragraph.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        fldChar_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._element.append(fldChar_begin)

        run2 = paragraph.add_run()
        run2.font.size = Pt(8)
        instrText = run2._element.makeelement(qn("w:instrText"), {})
        instrText.text = " PAGE "
        run2._element.append(instrText)

        run3 = paragraph.add_run()
        run3.font.size = Pt(8)
        fldChar_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run3._element.append(fldChar_end)

    def _maybe_add_table(self, doc: Document, profile: dict) -> None:
        """With ~38% probability, append a data table to the document."""
        if random.random() > 0.38:
            return

        table_type = random.choice(["budget", "status", "comparison", "contacts"])

        if table_type == "budget":
            doc.add_heading("Budget Summary", level=2)
            headers = ["Category", "Q1", "Q2", "Q3", "Q4"]
            rows = []
            for _ in range(random.randint(4, 7)):
                rows.append([
                    fake.bs().title(),
                    f"${random.randint(10, 500):,},000",
                    f"${random.randint(10, 500):,},000",
                    f"${random.randint(10, 500):,},000",
                    f"${random.randint(10, 500):,},000",
                ])
        elif table_type == "status":
            doc.add_heading("Project Status", level=2)
            headers = ["Task", "Owner", "Status", "Due Date"]
            statuses = ["Complete", "In Progress", "Not Started", "On Hold", "At Risk"]
            rows = []
            for _ in range(random.randint(5, 8)):
                rows.append([
                    fake.catch_phrase(),
                    fake.name(),
                    random.choice(statuses),
                    fake.date_between(start_date="+1d", end_date="+90d").strftime("%Y-%m-%d"),
                ])
        elif table_type == "comparison":
            doc.add_heading("Feature Comparison", level=2)
            headers = ["Feature", "Option A", "Option B", "Option C"]
            rows = []
            for _ in range(random.randint(5, 8)):
                rows.append([
                    fake.catch_phrase(),
                    random.choice(["Yes", "No", "Partial", "N/A"]),
                    random.choice(["Yes", "No", "Partial", "N/A"]),
                    random.choice(["Yes", "No", "Partial", "N/A"]),
                ])
        else:  # contacts
            doc.add_heading("Contact Directory", level=2)
            headers = ["Name", "Role", "Department", "Email"]
            departments = ["Engineering", "Sales", "Marketing", "Finance", "HR", "Operations", "Legal"]
            rows = []
            for _ in range(random.randint(4, 7)):
                rows.append([
                    fake.name(),
                    fake.job(),
                    random.choice(departments),
                    fake.company_email(),
                ])

        # Build the table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        try:
            table.style = profile["table_style"]
        except KeyError:
            table.style = "Table Grid"

        # Header row
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        # Data rows
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_val in enumerate(row_data):
                table.rows[r_idx + 1].cells[c_idx].text = cell_val

    def _generate_content(self, doc_type: str, context: Optional[str] = None) -> str:
        """Generate document content using LLM or fallback templates."""
        if self.llm:
            prompt = f"Generate a realistic {doc_type} document"
            if self.topic:
                prompt += f" related to {self.topic}"
            if context:
                prompt += f". Context from related email thread: {context}"
            prompt += """. 

CRITICAL RULES - This is a standalone business document (Word/PDF attachment), NOT AN EMAIL:
- NO email headers or metadata (no "Date:", "From:", "To:", "Subject:", "Re:", etc.)
- NO email greetings or signatures ("Dear", "Hi", "Best regards", "Sincerely", etc.)
- NO date headers at the top of the document
- NO "Prepared by" or "Prepared for" lines
- NO document metadata blocks

FORMAT: Start directly with the document title/heading, then the content.
Use markdown formatting for structure:
- Use ## for section headings and ### for sub-headings
- Use - or * for bullet lists
- Use 1. 2. 3. for numbered lists
- Use **bold** for key terms and emphasis
Mix these elements for a rich, professional layout.
Keep it under 750 words. Write ONLY the document content."""

            content = self.llm.generate_email_content(prompt)
            if content:
                return content

        # Fallback to topic-based template with markdown structure
        if self.topic:
            templates = [
                f"## Overview\n\n{fake.paragraph(nb_sentences=4)}\n\n"
                f"## Key Points\n\n"
                f"- {fake.sentence()}\n- {fake.sentence()}\n- {fake.sentence()}\n\n"
                f"## Details\n\n{fake.paragraph(nb_sentences=6)}",

                f"## Executive Summary\n\n{fake.paragraph(nb_sentences=5)}\n\n"
                f"## Analysis of {self.topic}\n\n{fake.paragraph(nb_sentences=4)}\n\n"
                f"### Findings\n\n"
                f"1. {fake.sentence()}\n2. {fake.sentence()}\n3. {fake.sentence()}\n\n"
                f"## Conclusion\n\n{fake.paragraph(nb_sentences=3)}",

                f"## {self.topic} Discussion Notes\n\n"
                f"**Attendees:** {fake.name()}, {fake.name()}, {fake.name()}\n\n"
                f"### Agenda Items\n\n"
                f"- {fake.sentence()}\n- {fake.sentence()}\n\n"
                f"### Action Items\n\n"
                f"1. {fake.sentence()}\n2. {fake.sentence()}\n\n"
                f"## Summary\n\n{fake.paragraph(nb_sentences=4)}",

                f"## Proposal: {self.topic}\n\n"
                f"### Background\n\n{fake.paragraph(nb_sentences=4)}\n\n"
                f"### Recommendations\n\n"
                f"- **Option A:** {fake.sentence()}\n"
                f"- **Option B:** {fake.sentence()}\n\n"
                f"### Next Steps\n\n{fake.paragraph(nb_sentences=3)}",
            ]
            return random.choice(templates)

        # Structured fallback instead of random faker text
        return (
            f"## Business Document\n\n"
            f"{fake.paragraph(nb_sentences=4)}\n\n"
            f"## Key Points\n\n"
            f"- {fake.sentence()}\n"
            f"- {fake.sentence()}\n"
            f"- {fake.sentence()}\n\n"
            f"## Summary\n\n"
            f"{fake.paragraph(nb_sentences=3)}"
        )

    def _generate_document_title(self, doc_type: str, context: Optional[str] = None) -> str:
        """Generate a professional document title using LLM or fallback."""
        if self.llm:
            prompt = f"""Generate a short, professional document filename (no extension) for a {doc_type}.
Context: {self.topic if self.topic else 'general business'}
Rules:
- Use 2-5 words maximum
- Use Title_Case_With_Underscores
- No dates, no special characters, no spaces
- Examples: Quarterly_Budget_Analysis, Project_Proposal, Meeting_Notes, Vendor_Agreement
Return ONLY the filename, nothing else."""
            title = self.llm.generate_email_content(prompt)
            if title:
                # Clean up any extra whitespace or quotes
                title = title.strip().strip('"').strip("'").strip()
                # Ensure proper formatting
                title = "_".join(word.capitalize() for word in title.replace("_", " ").split())
                title = sanitize_filename(title)
                if title:
                    return title
        
        # Fallback: combine topic and doc_type
        if self.topic:
            # Take first 2-3 words of topic
            words = self.topic.split()[:3]
            topic_part = "_".join(w.capitalize() for w in words)
            topic_part = sanitize_filename(topic_part)
            doc_type_cap = doc_type.capitalize()
            return f"{topic_part}_{doc_type_cap}"
        
        # Ultimate fallback
        return f"Business_{doc_type.capitalize()}"

    def generate_random_file(self, doc_type: str = "document", context: Optional[str] = None) -> str:
        """Generate a random PDF or DOCX file with LLM or fallback content."""
        ext = random.choice(["pdf", "docx"])
        # Generate a clean, professional document title
        doc_title = self._generate_document_title(doc_type, context)
        filename = f"{doc_title}.{ext}"
        content = self._generate_content(doc_type, context)

        if ext == "pdf":
            return self.create_pdf(filename, content)
        else:
            return self.create_docx(filename, content)
